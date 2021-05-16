import docx
import joblib
from __Parameters import *


def file_remove_enter():
    file_out = docx.Document()
    file_in = docx.Document(DATA)
    for i in file_in.paragraphs:
        if len(i.text) == 0:
            continue
        file_out.add_paragraph(i.text)
    file_out.save(DATA_REMOVED_ENTER)


def build_patient_locus():
    patient_idx_locus = {}
    file = docx.Document(DATA_REMOVED_ENTER)
    for i in range(0, len(file.paragraphs) - 1, 2):
        if file.paragraphs[i] not in patient_idx_locus:
            patient_idx = file.paragraphs[i].text[2:]
            patient_idx_locus[patient_idx] = file.paragraphs[i + 1].text
            print('patient_idx', patient_idx, 'finished')
    joblib.dump(patient_idx_locus, './patient/' + 'patient_idx_locus.pkl')

        
        
        
    

        